import discord
from discord.ext import commands
from discord import Embed
import openpyxl
import os
from dotenv import load_dotenv
from docx import Document

load_dotenv('.env')

def get_data_from_word_double(path_to_file):
    # Check if the file exists
    if not os.path.exists(path_to_file):
        raise FileNotFoundError(f"The file {path_to_file} does not exist.")
    
    # Check if the file is a docx file
    if not path_to_file.endswith('.docx'):
        raise ValueError(f"The file {path_to_file} is not a docx file.")
    
    #Creating a word file object
    doc_object = open(path_to_file, "rb")
            
    #creating word reader object
    doc_reader = Document(doc_object)
    data = ""
            
    for p in doc_reader.paragraphs:
        data += p.text+"\n\n"
                
    return data

def get_data_from_word_single(path_to_file):
    # Check if the file exists
    if not os.path.exists(path_to_file):
        raise FileNotFoundError(f"The file {path_to_file} does not exist.")
    
    # Check if the file is a docx file
    if not path_to_file.endswith('.docx'):
        raise ValueError(f"The file {path_to_file} is not a docx file.")
    
    #Creating a word file object
    doc_object = open(path_to_file, "rb")
            
    #creating word reader object
    doc_reader = Document(doc_object)
    data = ""
            
    for p in doc_reader.paragraphs:
        data += p.text+"\n"
                
    return data

def run_discord_bot():

    intents = discord.Intents.all()  

    bot = commands.Bot(command_prefix='!', intents=intents)

    token: str = os.getenv("DISCORD_TOKEN")

    """ CHANNELS """

    # RULES CHANNEL 
    rules_channel = int(os.getenv("RULES_CHANNEL"))

    # GENERAL INFORMATION CHANNELS
    important_information_channel = int(os.getenv("IMPORTANT_INFORMATION_CHANNEL"))

    # EXECUTIVE BRANCH CHANNELS
    exec_branch_information_channel = int(os.getenv("EXEC_BRANCH_CHANNEL"))
    nova_information_channel = int(os.getenv("NOVA_INFORMATION_CHANNEL"))
    internal_affairs_information_channel = int(os.getenv("INTERNAL_AFFAIRS_INFORMATION_CHANNEL"))
    external_affairs_information_channel = int(os.getenv("EXTERNAL_AFFAIRS_INFORMATION_CHANNEL"))
    student_media_information_channel = int(os.getenv("STUDENT_MEDIA_INFORMATION_CHANNEL"))

    # LEGISLATIVE BRANCH CHANNELS
    legislative_branch_information_channel = int(os.getenv("LEGISLATIVE_BRANCH_CHANNEL"))  
    lec_information_channel  = int(os.getenv("LEC_INFORMATION_CHANNEL"))
    abc_information_channel = int(os.getenv("ABC_INFORMATION_CHANNEL"))
    acc_information_channel = int(os.getenv("ACC_INFORMATION_CHANNEL"))
    pec_information_channel = int(os.getenv("PEC_INFORMATION_CHANNEL"))

    # JUDICIAL BRANCH CHANNELS
    judicial_branch_information_channel = int(os.getenv("JUDICIAL_BRANCH_CHANNEL"))
    judicial_information_channel = int(os.getenv("JUDICIAL_INFORMATION_CHANNEL"))

    """ FILE PATHS """

    # RULES
    rules_document = get_data_from_word_single("files/rules.docx")

    # MEMBER LIST
    workbook = openpyxl.load_workbook('files/member_list.xlsx')

    # EXECUTIVE BRANCH DOCUMENTS
    exec_information_document = get_data_from_word_double("files/exec/exec_info.docx")
    nova_information_document = get_data_from_word_double("files/exec/nova_info.docx")
    internal_affairs_information_document = get_data_from_word_double("files/exec/internal_affairs_info.docx")
    external_affairs_information_document = get_data_from_word_double("files/exec/external_affairs_info.docx")
    student_media_information_document = get_data_from_word_double("files/exec/student_media_info.docx")

    # LEGISLATIVE BRANCH DOCUMENTS
    legislative_branch_information_document = get_data_from_word_double("files/legislative/legislative_info.docx")
    lec_information_document = get_data_from_word_double("files/legislative/lec_info.docx")
    abc_information_document = get_data_from_word_double("files/legislative/abc_info.docx")
    acc_information_document = get_data_from_word_double("files/legislative/acc_info.docx")
    pec_information_document = get_data_from_word_double("files/legislative/pec_info.docx")

    # JUDICIAL BRANCH DOCUMENTS
    judicial_branch_information_document = get_data_from_word_double("files/judicial/judicial_branch_info.docx")
    judicial_information_document = get_data_from_word_double("files/judicial/judicial_info.docx")
    
    @bot.event
    async def on_ready():
        print(f'{bot.user} is now running!')

    # RULES COMMANDS
    @bot.command(name='rules')
    async def rules(ctx):
        channel = bot.get_channel(rules_channel)

        await channel.send(rules_document)

    """ POSITIONS COMMANDS """

    @bot.command(name='positions_exec_board')
    async def positions_exec_board(ctx):
        sheet = workbook.active
        
        channel = bot.get_channel(important_information_channel)
        
        branch = f"-----------------------------------------------------------\n**EXECUTIVE BOARD**\n-----------------------------------------------------------"
        target_sheet_name = "executive_board"

        if target_sheet_name in workbook.sheetnames:
            sheet = workbook[target_sheet_name]
            await channel.send(branch)

            for row in sheet.iter_rows(min_row=2, values_only=True):
                formatted_text = f"**{row[0]}:** {row[1]}"
                await channel.send(formatted_text)

    @bot.command(name='positions_exec_cabinet')
    async def positions_exec_cabinet(ctx):
        sheet = workbook.active
        
        channel = bot.get_channel(important_information_channel)
        
        branch = f"-----------------------------------------------------------\n**EXECUTIVE CABINET**\n-----------------------------------------------------------"
        target_sheet_name = "executive_cabinet"

        if target_sheet_name in workbook.sheetnames:
            sheet = workbook[target_sheet_name]
            await channel.send(branch)

            for row in sheet.iter_rows(min_row=2, values_only=True):
                formatted_text = f"**{row[0]}:** {row[1]}"
                await channel.send(formatted_text)

    @bot.command(name='positions_legislative')
    async def positions_legislative(ctx):
        sheet = workbook.active

        channel = bot.get_channel(important_information_channel)

        branch = f"-----------------------------------------------------------\n**LEGISLATIVE**\n-----------------------------------------------------------"
        target_sheet_name = "legislative"

        if target_sheet_name in workbook.sheetnames:
            sheet = workbook[target_sheet_name]
            await channel.send(branch)

            for row in sheet.iter_rows(min_row=2, values_only=True):
                formatted_text = f"**{row[0]}:** {row[1]}"
                await channel.send(formatted_text)

    @bot.command(name='positions_judicial')
    async def positions_judicial(ctx):
        sheet = workbook.active
        
        channel = bot.get_channel(important_information_channel)
        
        branch = f"-----------------------------------------------------------\n**JUDICIAL**\n-----------------------------------------------------------"
        target_sheet_name = "judicial"

        if target_sheet_name in workbook.sheetnames:
            sheet = workbook[target_sheet_name]
            await channel.send(branch)

            for row in sheet.iter_rows(min_row=2, values_only=True):
                formatted_text = f"**{row[0]}:** {row[1]}"
                await channel.send(formatted_text)

    
    """ INFORMATIONAL COMMANDS"""

    # EXECUTIVE BRANCH COMMANDS
    @bot.command(name='branch_info_exec')
    async def branch_info_exec(ctx):
        channel = bot.get_channel(exec_branch_information_channel)

        await channel.send(exec_information_document)

    @bot.command(name='nova_info')
    async def nova_info(ctx):
        channel = bot.get_channel(nova_information_channel)

        await channel.send(nova_information_document)

    @bot.command(name='internal_affairs_info')
    async def internal_affairs_info(ctx):
        channel = bot.get_channel(internal_affairs_information_channel)

        await channel.send(internal_affairs_information_document)

    @bot.command(name='external_affairs_info')
    async def external_affairs_info(ctx):
        channel = bot.get_channel(external_affairs_information_channel)

        await channel.send(external_affairs_information_document)

    @bot.command(name='student_media_info')
    async def student_media_info(ctx):
        channel = bot.get_channel(student_media_information_channel)

        await channel.send(student_media_information_document)

    # LEGISLATIVE BRANCH COMMANDS

    @bot.command(name='branch_info_legislative')
    async def branch_info_legislative(ctx):
        channel = bot.get_channel(legislative_branch_information_channel)

        await channel.send(legislative_branch_information_document)

    @bot.command(name='branch_info_lec')
    async def lec_info(ctx):
        channel = bot.get_channel(lec_information_channel)

        await channel.send(lec_information_document)

    @bot.command(name='abc_info')
    async def abc_info(ctx):
        channel = bot.get_channel(abc_information_channel)

        await channel.send(abc_information_document)

    @bot.command(name='acc_info')
    async def acc_info(ctx):
        channel = bot.get_channel(acc_information_channel)

        await channel.send(acc_information_document)

    @bot.command(name='pec_info')
    async def pec_info(ctx):
        channel = bot.get_channel(pec_information_channel)

        await channel.send(pec_information_document)

    # JUDICIAL BRANCH COMMANDS

    @bot.command(name='branch_info_judicial')
    async def branch_info_judicial(ctx):
        channel = bot.get_channel(judicial_branch_information_channel)

        await channel.send(judicial_branch_information_document)

    @bot.command(name='judicial_info')
    async def judicial_info(ctx):
        channel = bot.get_channel(judicial_information_channel)

        await channel.send(judicial_information_document)

    bot.run(token)
