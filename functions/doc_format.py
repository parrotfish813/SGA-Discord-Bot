from docx import Document
import os

def get_data_from_word_double(path_to_file):
    # Check if the file exists
    if not os.path.exists(path_to_file):
        raise FileNotFoundError(f"The file {path_to_file} does not exist.")
    
    # Check if the file is a docx file
    if not path_to_file.endswith('.docx'):
        raise ValueError(f"The file {path_to_file} is not a docx file.")
    
    #Creating a word file object
    doc_object = open(path_to_file, "rb")
            
    #creating word reader object
    doc_reader = Document(doc_object)
    data = ""
            
    for p in doc_reader.paragraphs:
        data += p.text+"\n\n"
                
    return data

def get_data_from_word_single(path_to_file):
    # Check if the file exists
    if not os.path.exists(path_to_file):
        raise FileNotFoundError(f"The file {path_to_file} does not exist.")
    
    # Check if the file is a docx file
    if not path_to_file.endswith('.docx'):
        raise ValueError(f"The file {path_to_file} is not a docx file.")
    
    #Creating a word file object
    doc_object = open(path_to_file, "rb")
            
    #creating word reader object
    doc_reader = Document(doc_object)
    data = ""
            
    for p in doc_reader.paragraphs:
        data += p.text+"\n"
                
    return data